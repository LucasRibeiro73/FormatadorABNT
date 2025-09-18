import docx
import re
from docx.shared import Cm, Pt, RGBColor # <-- ADICIONADO RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def formatador_abnt_profissional(caminho_entrada):
    print(f"Iniciando a formatação profissional do arquivo: {caminho_entrada}")
    doc = docx.Document(caminho_entrada)

    # --- 1. CONFIGURAÇÕES GLOBAIS E DE ESTILOS ---
    print("1/4 - Ajustando margens e redefinindo estilos ABNT...")
    for section in doc.sections:
        section.top_margin = Cm(3); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(2)
    
    # Redefinindo o estilo Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'; style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)

    # Redefinindo o estilo Título 1 (Seção Primária) para o padrão ABNT
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Arial'; style_h1.font.size = Pt(12)
    style_h1.font.bold = True; style_h1.font.uppercase = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    p_format_h1 = style_h1.paragraph_format
    p_format_h1.space_before = Pt(18); p_format_h1.space_after = Pt(6)
    p_format_h1.line_spacing = 1.5

    # Redefinindo o estilo Título 2 (Seção Secundária) para o padrão ABNT
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Arial'; style_h2.font.size = Pt(12)
    style_h2.font.bold = True; style_h2.font.uppercase = False
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    p_format_h2 = style_h2.paragraph_format
    p_format_h2.space_before = Pt(12); p_format_h2.space_after = Pt(6)
    p_format_h2.line_spacing = 1.5

    # --- 2. LÓGICA DE FORMATAÇÃO E APLICAÇÃO DE ESTILOS ---
    print("2/4 - Formatando o conteúdo e aplicando estilos...")
    in_textual_section = False; in_references_section = False
    in_abstract_section = False

    for paragrafo in doc.paragraphs:
        texto_paragrafo = paragrafo.text.strip()
        
        if 'ABSTRACT' in texto_paragrafo.upper(): in_abstract_section = True
        elif '1 INTRODUÇÃO' in texto_paragrafo.upper(): in_textual_section = True; in_abstract_section = False
        elif texto_paragrafo.upper() == 'REFERÊNCIAS': in_textual_section = False; in_references_section = True

        if in_abstract_section:
            if texto_paragrafo.upper() == 'ABSTRACT':
                paragrafo.style = doc.styles['Normal']
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                texto_original = paragrafo.text; paragrafo.text = ''
                paragrafo.add_run(texto_original).bold = True
                continue
            if texto_paragrafo:
                paragrafo.paragraph_format.line_spacing = 1.5; paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif in_textual_section:
            match_titulo = re.match(r'^(\d[\.\d]*)\s', texto_paragrafo)
            is_citacao_longa = texto_paragrafo.upper().startswith('[CITAÇÃO]')
            
            if match_titulo:
                numero_secao = match_titulo.group(1); nivel = numero_secao.count('.') + 1
                paragrafo.text = texto_paragrafo
                if nivel == 1: paragrafo.style = 'Heading 1'
                else: paragrafo.style = 'Heading 2'

            elif is_citacao_longa:
                paragrafo.text = texto_paragrafo.replace('[CITAÇÃO]', '', 1).strip(); paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragrafo.runs: run.font.size = Pt(10)
                p_format = paragrafo.paragraph_format; p_format.left_indent = Cm(4); p_format.first_line_indent = Cm(0); p_format.line_spacing = 1.0; p_format.space_before = Pt(6); p_format.space_after = Pt(6)

            elif texto_paragrafo != "":
                paragrafo.style = 'Normal'
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_format = paragrafo.paragraph_format; p_format.line_spacing = 1.5; p_format.first_line_indent = Cm(1.25); p_format.space_before = Pt(0); p_format.space_after = Pt(6)
        
        elif in_references_section:
            if texto_paragrafo.upper() == 'REFERÊNCIAS':
                paragrafo.style = 'Normal'
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                texto_original = paragrafo.text; paragrafo.text = ''
                paragrafo.add_run(texto_original).bold = True
                continue
            if texto_paragrafo:
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_format = paragrafo.paragraph_format; p_format.first_line_indent = Cm(0); p_format.line_spacing = 1.0; p_format.space_before = Pt(0); p_format.space_after = Pt(12)

    print("3/4 - Conteúdo formatado com estilos estruturais.")
    
    # --- 3. SALVANDO O ARQUIVO FINAL ---
    print("4/4 - Salvando o arquivo final...")
    caminho_saida = caminho_entrada.replace('.docx', '_FORMATADO_PROFISSIONAL.docx')
    doc.save(caminho_saida)
    print(f"\nPROCESSO CONCLUÍDO! Arquivo salvo em: {caminho_saida}")

if __name__ == '__main__':
    arquivo_do_aluno = 'meu_trabalho.docx'
    formatador_abnt_profissional(arquivo_do_aluno)